from telegram import Update, InputFile
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, ContextTypes, filters
from googletrans import Translator
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import docx
import os

translator = Translator()

# ====== START ======
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "👋 أهلاً بك في بوت ترجمة الملازم\n\n"
        "📌 يدعم:\n"
        "- PDF نصي\n"
        "- PDF مصوَّر (صور)\n"
        "- Word (DOCX)\n\n"
        "📘 الترجمة: إنكليزي ➜ عربي\n"
        "📗 الطريقة: سطر فوق سطر\n\n"
        "📤 فقط أرسل الملف مباشرة"
    )

# ====== HELP ======
async def help_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "🛠 طريقة الاستخدام:\n"
        "1️⃣ أرسل ملف PDF أو Word\n"
        "2️⃣ انتظر الترجمة\n"
        "3️⃣ استلم الملف مترجم\n\n"
        "⚠️ ملاحظة:\n"
        "- PDF المصوّر أبطأ شوي\n"
        "- الملفات النصية أدق"
    )

# ====== TRANSLATE TEXT ======
def translate_lines(text):
    lines = text.split("\n")
    result = ""
    for line in lines:
        if line.strip():
            ar = translator.translate(line, src="en", dest="ar").text
            result += ar + "\n" + line + "\n\n"
    return result

# ====== PDF HANDLER ======
async def handle_pdf(update: Update, context: ContextTypes.DEFAULT_TYPE):
    file = await update.message.document.get_file()
    input_pdf = "input.pdf"
    output_pdf = "translated.pdf"

    await file.download_to_drive(input_pdf)

    doc = fitz.open(input_pdf)
    new_doc = fitz.open()

    for page in doc:
        text = page.get_text("text")

        if not text.strip():
            # OCR
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = pytesseract.image_to_string(img)

        translated = translate_lines(text)

        new_page = new_doc.new_page()
        new_page.insert_textbox(
            new_page.rect,
            translated,
            fontsize=10
        )

    new_doc.save(output_pdf)
    doc.close()
    new_doc.close()

    await update.message.reply_document(InputFile(output_pdf))

    os.remove(input_pdf)
    os.remove(output_pdf)

# ====== WORD HANDLER ======
async def handle_docx(update: Update, context: ContextTypes.DEFAULT_TYPE):
    file = await update.message.document.get_file()
    input_docx = "input.docx"
    output_docx = "translated.docx"

    await file.download_to_drive(input_docx)

    doc = docx.Document(input_docx)
    new_doc = docx.Document()

    for para in doc.paragraphs:
        if para.text.strip():
            ar = translator.translate(para.text, src="en", dest="ar").text
            new_doc.add_paragraph(ar)
            new_doc.add_paragraph(para.text)
            new_doc.add_paragraph("")

    new_doc.save(output_docx)

    await update.message.reply_document(InputFile(output_docx))

    os.remove(input_docx)
    os.remove(output_docx)

# ====== MAIN ======
app = ApplicationBuilder().token("8537568402:AAHXW0gSYoBeZCIKokWgWXJGqKpg5mKj4N8").build()

app.add_handler(CommandHandler("start", start))
app.add_handler(CommandHandler("help", help_cmd))
app.add_handler(MessageHandler(filters.Document.PDF, handle_pdf))
app.add_handler(MessageHandler(filters.Document.DOCX, handle_docx))

print("Bot is running...")
app.run_polling()
